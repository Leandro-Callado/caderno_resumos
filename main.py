import os
import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Caminhos
TEMPLATE_PATH = "template_caderno.docx"
RESUMOS_DIR = "dados/resumos"
SAIDA = "Caderno_de_Resumos.docx"

def extrair_texto_pdf(caminho_pdf):
    """Extrai texto bruto de um PDF."""
    texto = ""
    with pdfplumber.open(caminho_pdf) as pdf:
        for pagina in pdf.pages:
            texto += pagina.extract_text() + "\n"
    return texto.strip()

def processar_resumo(texto):
    """
    Separa TÍTULO, AUTORES e CONTEÚDO.
    Assume:
        Linha 1 = Título
        Linha 2 = Autores
        Resto = conteúdo
    """
    linhas = texto.splitlines()
    linhas = [l.strip() for l in linhas if l.strip()]

    if len(linhas) < 3:
        raise ValueError("Formato inválido de resumo.")

    titulo = linhas[0]
    autores = linhas[1]
    conteudo = "\n".join(linhas[2:])

    return titulo, autores, conteudo


def adicionar_resumo(doc, titulo, autores, conteudo):
    """Insere um resumo já estruturado dentro do documento."""

    # ---- TÍTULO (Heading 2) ----
    h2 = doc.add_heading(titulo, level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.runs[0]
    run.font.name = "Libre Baskerville"
    run.font.size = Pt(14)
    run.bold = True

    # ---- Autores ----
    p_aut = doc.add_paragraph(autores)
    p_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_aut.runs[0]
    run.font.name = "Libre Baskerville"
    run.font.size = Pt(11)

    doc.add_paragraph("")  # espaço

    # ---- Conteúdo ----
    for linha in conteudo.split("\n"):
        p = doc.add_paragraph(linha)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Libre Baskervile"
            run.font.size = Pt(11)

    doc.add_page_break()


def gerar_caderno():
    print("📘 Gerando Caderno completo...")

    # Carregar template
    doc = Document(TEMPLATE_PATH)

    # Procurar PDFs
    pdfs = sorted([
        f for f in os.listdir(RESUMOS_DIR)
        if f.lower().endswith(".pdf")
    ])

    if not pdfs:
        print("⚠ Nenhum PDF encontrado em dados/resumos/")
        return

    for nome_pdf in pdfs:
        caminho = os.path.join(RESUMOS_DIR, nome_pdf)
        print(f"→ Importando: {nome_pdf}")

        texto = extrair_texto_pdf(caminho)
        titulo, autores, conteudo = processar_resumo(texto)
        adicionar_resumo(doc, titulo, autores, conteudo)

    # Salvar documento final
    doc.save(SAIDA)
    print(f"✅ Caderno gerado com sucesso: {SAIDA}")


if __name__ == "__main__":
    gerar_caderno()
